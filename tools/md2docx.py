# -*- coding: utf-8 -*-
"""md2docx.py — 项目文档约定工具：把 Markdown 交付物转成自然语言 Word 版
用法: python md2docx.py <file1.md> [file2.md ...]
输出: 同目录同名 .docx（微软雅黑、表格转真表格、无 markdown 符号）
"""
import re
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

CN_FONT = "微软雅黑"
MONO_FONT = "Consolas"
HEAD_COLOR = (0x1B, 0x5E, 0x5A)  # Wellcee 青绿


def set_font(run, name=CN_FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def clean_inline(text):
    """去掉链接语法，保留可读文本"""
    return LINK_RE.sub(lambda m: f"{m.group(1)}（{m.group(2)}）", text)


def add_inline(par, text, base_bold=False, size=None):
    """解析 **bold** 与 `code`，写入 runs"""
    text = clean_inline(text)
    for seg in INLINE_RE.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**") and len(seg) > 4:
            r = par.add_run(seg[2:-2])
            set_font(r, bold=True, size=size)
        elif seg.startswith("`") and seg.endswith("`") and len(seg) > 2:
            r = par.add_run(seg[1:-1])
            set_font(r, name=MONO_FONT, size=size)
        else:
            r = par.add_run(seg)
            set_font(r, bold=base_bold or None, size=size)


def strip_md(text):
    """单元格等场景：去掉全部行内标记，返回纯文本"""
    text = clean_inline(text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def flush_table(doc, rows):
    if not rows:
        return
    # 过滤分隔行 |---|---|
    body = [r for r in rows if not re.fullmatch(r"[\s|:\-]+", r)]
    parsed = []
    for r in body:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(parsed):
        for j in range(ncols):
            cell = table.cell(i, j)
            txt = strip_md(row[j]) if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_font(r, size=9, bold=(i == 0))
    doc.add_paragraph()


def convert(md_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    table_buf = []
    in_code = False
    for raw in lines:
        line = raw.rstrip("\n")
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            r = p.add_run(line)
            set_font(r, name=MONO_FONT, size=9)
            continue

        # 表格行缓冲
        if stripped.startswith("|"):
            table_buf.append(stripped)
            continue
        elif table_buf:
            flush_table(doc, table_buf)
            table_buf = []

        if not stripped:
            continue
        if re.fullmatch(r"-{3,}", stripped):
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            sizes = {1: 18, 2: 14, 3: 12, 4: 11}
            p = doc.add_heading("", level=level)
            add_inline(p, strip_md(m.group(2)), size=sizes[level])
            for r in p.runs:
                set_font(r, size=sizes[level], bold=True,
                         color=HEAD_COLOR if level <= 2 else (0x33, 0x33, 0x33))
            continue

        # 图片：![alt](path) 独占一行 → 嵌入图（相对路径按 md 所在目录解析）
        m = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            img = Path(m.group(2))
            if not img.is_absolute():
                img = md_path.parent / img
            if img.exists():
                try:
                    doc.add_picture(str(img), width=Cm(16))
                except Exception:
                    p = doc.add_paragraph()
                    r = p.add_run(f"[图：{m.group(1) or img.name}]")
                    set_font(r, italic=True, color=(0x88, 0x88, 0x88))
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"[缺图：{m.group(2)}]")
                set_font(r, italic=True, color=(0xAA, 0x33, 0x33))
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            r = p.add_run(strip_md(stripped.lstrip("> ")))
            set_font(r, italic=True, size=10, color=(0x55, 0x55, 0x55))
            continue

        # checkbox
        m = re.match(r"^-\s+\[( |x|X)\]\s+(.*)", stripped)
        if m:
            mark = "☑ " if m.group(1).lower() == "x" else "☐ "
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(mark)
            set_font(r)
            add_inline(p, m.group(2))
            continue

        # 无序列表（两级）
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, m.group(2))
            continue

        # 有序列表：保留编号文本
        m = re.match(r"^(\s*)(\d+)[.、]\s+(.*)", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75 if len(m.group(1)) < 2 else 1.5)
            r = p.add_run(f"{m.group(2)}. ")
            set_font(r, bold=True)
            add_inline(p, m.group(3))
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_inline(p, stripped)

    if table_buf:
        flush_table(doc, table_buf)

    out = md_path.with_suffix(".docx")
    doc.save(out)
    print(f"OK  {md_path.name}  ->  {out.name}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python md2docx.py <file.md> [more.md ...]")
        sys.exit(1)
    for arg in sys.argv[1:]:
        convert(Path(arg))
